from flask import Flask, render_template, request, jsonify, send_file
from bs4 import BeautifulSoup
import requests
from sklearn.feature_extraction.text import TfidfVectorizer
from googletrans import Translator

from docx import Document
from datetime import datetime
import time
import ftplib
import time


# 뉴스 페이지를 순회하며 뉴스의 링크를 수집
def collect_links():
    links = []
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.6367.118 Safari/537.36"
    }
    for i in range(1, 11):
        outter_url = f"http://www.dailysecu.com/news/articleList.html?page={i}&total=13290&box_idxno=&sc_section_code=S1N2&view_type=sm"
        outter_req = requests.get(outter_url, headers=headers)
        outter_soup = BeautifulSoup(outter_req.text, "lxml")
        outter_tags = outter_soup.select(
            "#user-container > div > div > section > article > div > section > div > div > a"
        )
        links += [f"http://www.dailysecu.com{tag.get('href')}" for tag in outter_tags]
    return links


# 뉴스의 본문을 수집
def extract_descriptions(links):
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.6367.118 Safari/537.36"
    }
    articles = []
    for link in links:
        inner_req = requests.get(link, headers=headers)
        inner_soup = BeautifulSoup(inner_req.text, "lxml")
        inner_tags = inner_soup.select("#user-container > div > div > section > div p")
        p = ""
        for p_tag in inner_tags:
            if p_tag.text.startswith("▷") or p_tag.text.startswith("★"):
                continue
            p += p_tag.text.strip()
        articles.append(str(p))

    return articles


# tf-idf를 위해 내용을 영어로 번역함
def translate_for_tfidf(articles):
    translator = Translator(
        service_urls=[
            "translate.google.com",
            "translate.google.co.kr",
        ]
    )

    i = 0
    for index, article in enumerate(articles):
        print(i)
        i += 1
        try:
            # 언어 감지
            detected = translator.detect(article)
            if detected is not None:
                # 텍스트 번역
                translated = translator.translate(article, src=detected.lang, dest="en")
                if translated is not None:
                    articles[index] = translated.text
                else:
                    print("Translation failed: No response from API.")
            else:
                print("Language detection failed: No response from API.")

        except Exception as e:
            print(f"An error occurred: {e}")

    return articles


# 단어의 중요도를 평가하기위한 tf-idf
def cal_tfidf(articles):
    # TF-IDF 벡터화 객체 생성
    vectorizer = TfidfVectorizer(stop_words="english")

    # TF-IDF 행렬 계산
    tfidf_matrix = vectorizer.fit_transform(articles)

    # 단어 목록
    terms = vectorizer.get_feature_names_out()

    # 각 단어의 TF-IDF 점수 계산
    tfidf_scores = tfidf_matrix.max(0).toarray()[0]

    # 단어와 TF-IDF 점수를 튜플 리스트로 저장
    word_tfidf_tuples = [(term, score) for term, score in zip(terms, tfidf_scores)]

    # TF-IDF 점수 기준으로 내림차순 정렬
    word_tfidf_tuples.sort(key=lambda x: x[1], reverse=True)

    return word_tfidf_tuples


# 상위 단어 40개 출력
def print_tfidf(word_tfidf_tuples):
    print("{:<15} {}".format("단어", "TF-IDF 점수"))
    print("=" * 30)
    for term, score in word_tfidf_tuples[:40]:
        print("{:<15} {:.2f}".format(term, score))


def create_report(word_tfidf_tuples):
    try:
        """보고서 생성하는 방식활용, 템플릿X"""
        # print(word_tfidf_tuples)
        # 워드 보고서 생성
        doc = Document()
        doc.add_heading("Latest Keywords of Secure Infomation ", level=0)
        # 날짜 데이터 가져오기
        now = datetime.now()
        day = now.strftime("%Y-%m-%d")
        doc.add_paragraph(day)

        """"-------------------워드에 테이블 삽입---------------------"""
        length = 40  # 상위 몇개의 단어를 가지고 오고 싶은지...?

        table = doc.add_table(rows=length + 1, cols=2)
        table.style = doc.styles["Table Grid"]
        header = table.rows[0].cells
        header[0].text = "Keywords"
        header[1].text = "Frequency"

        for i in range(1, length + 1):
            r = table.rows[i].cells
            r[0].text = word_tfidf_tuples[i - 1][0]
            r[1].text = f"{word_tfidf_tuples[i-1][1]:.2f}"

        doc.save("secure_news_keywords_report.docx")
        hostname = "192.168.122.128"
        ftp = ftplib.FTP(hostname)

        ftp.login("msfadmin", "msfadmin")

        with open("secure_news_keywords_report.docx", "rb") as file:
            ftp.storbinary("STOR secure_news_keywords_report.docx", file)
            print("보고서 백업 완료")
        ftp.quit()
        return True

    except Exception as e:
        print(f"실패 이유 {e}")
        return False


def main():
    app.run(debug=True)


app = Flask(__name__)


@app.route("/")
def home():
    return render_template("index.html")


@app.route("/create_secure_trend", methods=["GET", "POST"])
def create_secure_trend():

    task = request.form.get("task")
    links = collect_links()
    articles = extract_descriptions(links)
    translated = translate_for_tfidf(articles)
    tfidf = cal_tfidf(translated)
    print_tfidf(tfidf)

    keyword_report = create_report(tfidf)
    booan = "secure_news_keywords_report.docx"
    if task == "report":
        success = keyword_report
        if success:
            return send_file(booan, as_attachment=True)


def main():
    app.run(debug=True, host="0.0.0.0", port=5000)
    create_secure_trend()
    create_keyword_report()


if __name__ == "__main__":
    main()
